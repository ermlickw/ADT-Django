from docx import Document
from docx.shared import Inches

refA = "Smith"
refB="Jones"

class Claim:
    def __init__(self, references=refA, dependentOn=0,
                 number=1, text=" no text", citedText="no cited text", para=None):
        self.references = references
        self.dependentOn = dependentOn
        self.number = number
        self.text = text
        self.citedText = citedText
        self.para = para

    def write_claim(self):
        self.para = document.add_paragraph(self.citedText)
        self.para.add_run(self.references)

        # document.add_run('bold').bold = True
        # document.add_run(' and some ')
    def add_stuff(self,textstuff):
        self.para.add_run(textstuff)





# if __name__ == "__main__":

document = Document()

document.add_heading('Document Title', 0)

clm1 = Claim(citedText="stuff.", )
clm1.write_claim()
clm1.add_stuff("extra stuff")
print(clm1.para.text)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
q = document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
    'first item in ordered list', style='ListNumber'
)

document.save('demo.docx')
