from __future__ import print_function
from django.db import models
from django.utils import timezone
from django.urls import reverse
from django.conf import settings
from django.core.files import File
from docx import Document
import traceback
import sys
from .validators import validate_file_extension
from django.contrib.auth.models import User
# from django.core.files.storage import FileSystemStorage
#
# good place to start project is with models

class File(models.Model):

    def only_filename(instance,filename):
        return instance.appNumber + '.docx'

    appNumber = models.TextField(max_length=10)
    document = models.FileField(upload_to=only_filename, blank=True, validators=[validate_file_extension])
    created_at = models.DateTimeField(auto_now_add=True)
    # updated_at = models.DateTimeField(null=True)
    # description = models.TextField(blank=True)

    def __str__(self):
        return self.appNumber

    def get_absolute_url(self): #after post go where? only happens when {{form}} is provided in template
        return reverse("claims",kwargs={'pk':self.appNumber})





class Claim(models.Model):
    appNumber = models.TextField(max_length=15,blank=True)
    references = models.TextField(max_length=15,blank=True)
    dependentOn = models.IntegerField(default = 0,blank=True)
    number = models.IntegerField(default=0,blank=True)
    text = models.TextField(max_length=2000,blank=True)
    citedText = models.TextField(max_length=2000,blank=True)
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(null=True)
    # created_by = models.ForeignKey(User, on_delete=models.CASCADE)
    # updated_by = models.ForeignKey(User, null=True, on_delete=models.CASCADE)

    # def write_claim(self):
    #     self.para = document.add_paragraph(self.citedText)
    #     self.para.add_run(self.references)
    #
    # def add_stuff(self,textstuff):
    #     self.para.add_run(textstuff)

    def __str__(self):
        return str(self.appNumber) + "-Claim-" + str(self.number)
